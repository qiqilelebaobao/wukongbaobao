from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

class Printer:
    def __init__(self) -> None:
        self.document = None
        self.output_name = ''

    @staticmethod
    def gen_txt_text_from_question(qs):
        page_line = 35
        outer = '\n'
        if (l:=len(qs)) >= page_line:
            for i, q in enumerate(qs):
                if i % 2 == 0:
                    outer += f'{q}\t\t\t\t'
                else:
                    outer += f'{q}\n\n\n'
        else:
            for i, q in enumerate(qs):
                outer += f'{q}\n\n\n'
        return outer

    def txt_output(self, qs):
        self.output_name = f'output/plus_{datetime.now().strftime("%y%m%d_%H%M%S")}.txt'
        with open(self.output_name, 'w') as f:
            print(Printer.gen_txt_text_from_question(qs), file = f)

    @staticmethod
    def set_head(head, text):
        head.paragraph_format.space_before = Pt(10)
        head.paragraph_format.space_after = Pt(10)
        head.paragraph_format.line_spacing = 1.0
        run = head.add_run(text)
        run.font.name = u'简宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'简宋') # type: ignore
    
    @staticmethod
    def set_graph(p, text):
        p.paragraph_format.line_spacing = 2.5
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') # type: ignore
        # run.font.bold = True

    @staticmethod
    def gen_doc_para_text_from_question(qs):
        outer = ''
        for i, q in enumerate(qs):
            if i % 2 == 0:
                if len(q) < 15:
                    outer += f'{q}\t\t\t\t'
                else:
                    outer += f'{q}\n'
            else:
                outer += f'{q}\n'
        return outer

    def doc_output(self, qs):
        self.document = Document()
        document = self.document
        
        h = document.add_heading('', 0)
        Printer.set_head(h, '数学测试')
        
        p = document.add_paragraph()
        Printer.set_graph(p, Printer.gen_doc_para_text_from_question(qs))

        self.output_name = f'output/plus_{datetime.now().strftime("%y%m%d_%H%M%S")}.docx'
        document.save(self.output_name)

    def output(self, qs, formatt='txt') -> None:
        if formatt == 'txt':
            return self.txt_output(qs)
        elif formatt == 'doc' or formatt == 'docx':
            return self.doc_output(qs)
        else:
            print('wrong position...')
